import os
import re
import io
import time
import warnings
from typing import List, Optional

import numpy as np
import pandas as pd
import streamlit as st


# ===============================
# Page setup
# ===============================
st.set_page_config(page_title="NLP Tools", page_icon="🧠", layout="wide")
st.markdown("<h1 style='text-align: center;'>🧠 NLP Tools</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center;'>A unified interface for Sentiment Analysis, Named Entity Recognition, Question Answering, Summarization, and Resume Screening.</p>", unsafe_allow_html=True)


# ===============================
# Utilities
# ===============================
def _decode_bytes(data: bytes) -> str:
    if not isinstance(data, (bytes, bytearray)):
        return str(data)
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return data.decode(enc, errors="ignore")
        except Exception:
            continue
    return ""


def read_upload_any(file) -> str:
    """Extract text from an uploaded file (txt, pdf, docx). Uses available parsers gracefully."""
    name = (getattr(file, "name", "uploaded") or "uploaded").lower()
    data = file.read()
    bio = io.BytesIO(data)
    try:
        file.seek(0)
    except Exception:
        pass

    if name.endswith(".txt"):
        return _decode_bytes(data)

    if name.endswith(".pdf"):
        # Try pdfplumber first, then PyPDF2
        text = ""
        try:
            import pdfplumber  # type: ignore
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for p in pdf.pages:
                    text += p.extract_text() or ""
        except Exception:
            try:
                from PyPDF2 import PdfReader  # type: ignore
                reader = PdfReader(bio)
                pages = []
                for pg in reader.pages:
                    try:
                        pages.append(pg.extract_text() or "")
                    except Exception:
                        pages.append("")
                text = "\n".join(pages)
            except Exception:
                text = ""
        return text

    if name.endswith(".docx"):
        try:
            from docx import Document  # type: ignore
            doc = Document(bio)
            return "\n".join([p.text for p in doc.paragraphs if p.text and p.text.strip()])
        except Exception:
            return ""

    return ""


def normalize_text(s: str) -> str:
    if not isinstance(s, str):
        s = str(s) if s is not None else ""
    return re.sub(r"\s+", " ", s.replace("\r", " ").replace("\n", " ")).strip()


def sent_split(text: str) -> List[str]:
    return [s.strip() for s in re.split(r"[\n\.!?]+", text) if s.strip()]


def apply_cache(func):
    """Compatibility wrapper across Streamlit versions."""
    if hasattr(st, "cache_resource"):
        return st.cache_resource(show_spinner=True)(func)
    return st.cache(allow_output_mutation=True)(func)


# ===============================
# Sidebar navigation
# ===============================
APP = st.sidebar.selectbox(
    "Select tool",
    [
        "Sentiment Analysis",
        "Named Entity Recognition",
        "Question Answering",
        "Summarizer",
        "Resume Screening",
    ],
)


# ===============================
# Sentiment Analysis Page
# ===============================
def sentiment_page():
    st.header("💬 Sentiment Analysis")
    st.write("Load pre-trained scikit-learn models and analyze sentiment for text, CSV, or files.")

    # Load artifacts lazily
    @apply_cache
    def load_artifacts():
        try:
            import joblib  # type: ignore
        except Exception:
            return None, None, None, "joblib not installed. Install with: pip install joblib"

        arts = os.path.join(os.path.dirname(__file__), "sentiment_artifacts")
        vec_p = os.path.join(arts, "vectorizer.joblib")
        lr_p = os.path.join(arts, "logreg.joblib")
        nb_p = os.path.join(arts, "naive_bayes.joblib")
        if not (os.path.exists(vec_p) and os.path.exists(lr_p) and os.path.exists(nb_p)):
            return None, None, None, f"Artifacts not found in {arts}."
        try:
            vectorizer = joblib.load(vec_p)
            lr = joblib.load(lr_p)
            nb = joblib.load(nb_p)
            return vectorizer, lr, nb, None
        except Exception as e:
            return None, None, None, f"Failed to load artifacts: {e}"

    vectorizer, lr, nb, err = load_artifacts()
    if err:
        st.error(err)
        st.info("Ensure the training notebook saved artifacts to 'sentiment_artifacts/'.")

    # Cleaning: light-weight, robust
    PUNCT = str.maketrans('', '', __import__('string').punctuation)
    URL_RE = re.compile(r"https?://\S+|www\.\S+")
    HTML_RE = re.compile(r"<.*?>")

    @apply_cache
    def init_cleaners():
        # Try spaCy, then NLTK, else None
        nlp = None
        wnl = None
        stop_words = set()
        # spaCy
        try:
            import spacy  # type: ignore
            try:
                nlp = spacy.load("en_core_web_sm", disable=["parser", "ner"])  # fast
            except Exception:
                nlp = spacy.blank("en")
        except Exception:
            nlp = None
        # NLTK
        try:
            import nltk  # type: ignore
            from nltk.corpus import stopwords  # type: ignore
            from nltk.stem import WordNetLemmatizer  # type: ignore
            try:
                nltk.data.find('corpora/stopwords')
            except LookupError:
                nltk.download('stopwords')
            try:
                nltk.data.find('corpora/wordnet')
            except LookupError:
                nltk.download('wordnet'); nltk.download('omw-1.4')
            stop_words = set(stopwords.words("english"))
            # keep negations
            for w in {"no", "nor", "not", "n't", "never"}:
                if w in stop_words:
                    stop_words.remove(w)
            wnl = WordNetLemmatizer()
        except Exception:
            pass
        return nlp, wnl, stop_words

    nlp, wnl, stop_words = init_cleaners()

    def lemmatize(tokens: List[str]) -> List[str]:
        if nlp is not None:
            doc = nlp(" ".join(tokens))
            return [t.lemma_ for t in doc]
        if wnl is not None:
            return [wnl.lemmatize(t) for t in tokens]
        return tokens

    def clean_text(text: str) -> str:
        if not isinstance(text, str):
            text = str(text)
        text = text.lower()
        text = URL_RE.sub(" ", text)
        text = HTML_RE.sub(" ", text)
        text = text.translate(PUNCT)
        text = re.sub(r"\d+", " ", text)
        toks = [t for t in text.split() if t.isalpha()]
        toks = lemmatize(toks)
        if stop_words:
            toks = [t for t in toks if t not in stop_words and len(t) > 1]
        return " ".join(toks)

    def predict_both(raw: str):
        if vectorizer is None:
            return 0.5, 0.5
        x = vectorizer.transform([clean_text(raw)])
        p_lr = float(lr.predict_proba(x)[0][1]) if lr is not None else 0.5
        p_nb = float(nb.predict_proba(x)[0][1]) if nb is not None else 0.5
        return p_lr, p_nb

    with st.sidebar:
        st.subheader("Sentiment Settings")
        compare = st.checkbox("Compare both models", value=True)
        model_choice = st.radio("Model", ["Logistic Regression", "Naive Bayes"], index=0, disabled=compare)
        threshold = st.slider("Decision threshold", 0.1, 0.9, 0.5, 0.05)

    tab1, tab2, tab3 = st.tabs(["Single Text", "Batch CSV", "Files (PDF/TXT/DOCX)"])

    with tab1:
        txt = st.text_area("Enter text", height=160, placeholder="Type or paste text…")
        if st.button("Predict Sentiment", key="sent_single_btn", type="primary"):
            if not txt.strip():
                st.warning("Please enter some text.")
            else:
                p_lr, p_nb = predict_both(txt)
                if compare:
                    c1, c2 = st.columns(2)
                    with c1:
                        st.subheader("Logistic Regression")
                        st.write(f"Positive: {p_lr:.3f} — Negative: {1-p_lr:.3f}")
                        st.progress(min(max(p_lr, 0.0), 1.0))
                        st.write("Label:", "Positive" if p_lr >= threshold else "Negative")
                    with c2:
                        st.subheader("Naive Bayes")
                        st.write(f"Positive: {p_nb:.3f} — Negative: {1-p_nb:.3f}")
                        st.progress(min(max(p_nb, 0.0), 1.0))
                        st.write("Label:", "Positive" if p_nb >= threshold else "Negative")
                else:
                    if model_choice == "Logistic Regression":
                        lbl = "Positive" if p_lr >= threshold else "Negative"
                        st.subheader(f"Prediction: {lbl}")
                        st.write(f"Probability Positive: {p_lr:.3f} — Negative: {1-p_lr:.3f}")
                        st.progress(min(max(p_lr, 0.0), 1.0))
                    else:
                        lbl = "Positive" if p_nb >= threshold else "Negative"
                        st.subheader(f"Prediction: {lbl}")
                        st.write(f"Probability Positive: {p_nb:.3f} — Negative: {1-p_nb:.3f}")
                        st.progress(min(max(p_nb, 0.0), 1.0))

    with tab2:
        st.write("Upload a CSV with a text column (e.g., 'review'/'text').")
        file = st.file_uploader("CSV file", type=["csv"], key="sent_csv")
        if file is not None:
            try:
                df = pd.read_csv(file)
            except Exception as e:
                st.error(f"Failed to read CSV: {e}")
                df = None
            if df is not None:
                cand_cols = [c for c in df.columns if c.lower() in ["review","text","review_text","content","body"]]
                if not cand_cols:
                    st.error("No review text column found. Add a 'review' or 'text' column.")
                else:
                    col = cand_cols[0]
                    cleaned = df[col].astype(str).apply(clean_text)
                    X = vectorizer.transform(cleaned) if vectorizer is not None else None
                    proba_lr = (lr.predict_proba(X)[:, 1] if (X is not None and lr is not None) else np.full(len(df), 0.5))
                    proba_nb = (nb.predict_proba(X)[:, 1] if (X is not None and nb is not None) else np.full(len(df), 0.5))
                    if compare:
                        out = df.copy()
                        out["LR_proba_pos"] = proba_lr
                        out["LR_label"] = np.where(out["LR_proba_pos"]>=threshold, "Positive", "Negative")
                        out["NB_proba_pos"] = proba_nb
                        out["NB_label"] = np.where(out["NB_proba_pos"]>=threshold, "Positive", "Negative")
                        st.write("Preview:")
                        st.dataframe(out.head(20), use_container_width=True)
                        st.write("Counts (LR):", out["LR_label"].value_counts())
                        st.write("Counts (NB):", out["NB_label"].value_counts())
                        st.download_button("Download predictions CSV", out.to_csv(index=False).encode("utf-8"), file_name="sentiment_predictions_compare.csv", mime="text/csv")
                    else:
                        proba = proba_lr if model_choice == "Logistic Regression" else proba_nb
                        pred = (proba >= threshold).astype(int)
                        out = df.copy()
                        out["pred_label"] = np.where(pred==1, "Positive", "Negative")
                        out["pred_proba_pos"] = proba
                        st.write("Preview:")
                        st.dataframe(out.head(20), use_container_width=True)
                        st.write("Counts:", out["pred_label"].value_counts())
                        st.download_button("Download predictions CSV", out.to_csv(index=False).encode("utf-8"), file_name="sentiment_predictions.csv", mime="text/csv")

    with tab3:
        st.write("Upload one or more files: PDF, TXT, or DOCX.")
        files = st.file_uploader("Files", type=["pdf","txt","docx"], accept_multiple_files=True, key="sent_files")
        if files:
            rows = []
            for f in files:
                name = f.name
                try:
                    text = read_upload_any(f)
                except Exception:
                    text = ""
                text = normalize_text(text)
                if not text:
                    continue
                p_lr, p_nb = predict_both(text)
                if compare:
                    rows.append({
                        "file": name,
                        "LR_label": ("Positive" if p_lr>=threshold else "Negative"),
                        "LR_proba_pos": p_lr,
                        "NB_label": ("Positive" if p_nb>=threshold else "Negative"),
                        "NB_proba_pos": p_nb,
                        "chars": len(text)
                    })
                else:
                    if model_choice == "Logistic Regression":
                        rows.append({"file": name, "label": ("Positive" if p_lr>=threshold else "Negative"), "proba_pos": p_lr, "chars": len(text)})
                    else:
                        rows.append({"file": name, "label": ("Positive" if p_nb>=threshold else "Negative"), "proba_pos": p_nb, "chars": len(text)})
            if rows:
                res = pd.DataFrame(rows)
                st.dataframe(res, use_container_width=True)
                st.download_button("Download results CSV", res.to_csv(index=False).encode("utf-8"), file_name="file_sentiment_predictions.csv", mime="text/csv")


# ===============================
# Summarizer Page
# ===============================
def summarizer_page():
    st.header("📝 Summarizer")
    st.write("Abstractive summarization for long texts with optional chunking.")

    CANDIDATE_MODELS = [
        "summarizer_ft\\final",
        "summarizer_ft\\checkpoint-50",
        "ainize/bart-base-cnn",
        "google-t5/t5-small",
    ]

    @apply_cache
    def load_model(name: str):
        try:
            import torch  # type: ignore
            from transformers import AutoTokenizer, AutoModelForSeq2SeqLM  # type: ignore
        except Exception as e:
            return None, None, f"Transformers not available: {e}"
        tok = AutoTokenizer.from_pretrained(name)
        mdl = AutoModelForSeq2SeqLM.from_pretrained(name)
        device = "cuda" if hasattr(torch, "cuda") and torch.cuda.is_available() else "cpu"
        mdl = mdl.to(device)
        mdl.eval()
        return (tok, mdl, device)

    # Sidebar controls
    model_name = st.sidebar.selectbox("Summarizer model", options=CANDIDATE_MODELS, index=0)
    max_input_tokens = st.sidebar.slider("Max input tokens", 256, 2048, 1024, 64)
    summary_max_tokens = st.sidebar.slider("Summary max tokens", 16, 256, 128, 8)
    num_beams = st.sidebar.slider("Beam size", 1, 6, 4, 1)
    use_chunking = st.sidebar.checkbox("Use chunking for long texts", value=True)

    def safe_sent_tokenize(text: str):
        return [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if s.strip()]

    def chunk_text(text: str, tokenizer, max_input_tokens: int = 1024, overlap_sentences: int = 1):
        sents = safe_sent_tokenize(text)
        chunks, buf = [], []
        for s in sents:
            candidate = " ".join(buf + [s])
            if len(tokenizer.encode(candidate, truncation=False)) <= max_input_tokens:
                buf.append(s)
            else:
                if buf:
                    chunks.append(" ".join(buf))
                    buf = buf[-overlap_sentences:] if overlap_sentences > 0 else []
                    buf.append(s)
                else:
                    tokens = tokenizer.encode(s, truncation=True, max_length=max_input_tokens)
                    chunks.append(tokenizer.decode(tokens, skip_special_tokens=True))
                    buf = []
        if buf:
            chunks.append(" ".join(buf))
        return chunks

    def summarize_chunks(chunks, tokenizer, model, device, summary_max_tokens=128, num_beams=4):
        import torch  # type: ignore
        parts = []
        for ch in chunks:
            inputs = tokenizer(ch, return_tensors="pt", truncation=True, max_length=min(2048, tokenizer.model_max_length)).to(device)
            with torch.no_grad():
                ids = model.generate(**inputs, max_new_tokens=summary_max_tokens, num_beams=num_beams, early_stopping=True)
            parts.append(tokenizer.decode(ids[0], skip_special_tokens=True))
        return " ".join(parts)

    def summarize(text: str, tokenizer, model, device, max_input_tokens=1024, summary_max_tokens=128, use_chunking=True, num_beams=4):
        is_t5 = ("t5" in model_name.lower())
        t = ("summarize: " + text) if is_t5 else text
        tokens = tokenizer.encode(t, truncation=False)
        if len(tokens) <= max_input_tokens or not use_chunking:
            inputs = tokenizer(t, return_tensors="pt", truncation=True, max_length=max_input_tokens).to(device)
            import torch  # type: ignore
            with torch.no_grad():
                ids = model.generate(**inputs, max_new_tokens=summary_max_tokens, num_beams=num_beams, early_stopping=True)
            return tokenizer.decode(ids[0], skip_special_tokens=True)
        chunks = chunk_text(t, tokenizer, max_input_tokens=max_input_tokens)
        return summarize_chunks(chunks, tokenizer, model, device, summary_max_tokens=summary_max_tokens, num_beams=num_beams)

    # File upload and input
    uploaded = st.file_uploader("Upload PDF, DOCX, or TXT", type=["pdf","docx","txt"], accept_multiple_files=True, key="sum_files")
    if uploaded:
        texts = []
        for f in uploaded:
            t = read_upload_any(f)
            if t:
                texts.append(t)
            else:
                st.warning(f"Couldn't extract text from {getattr(f,'name','file')}")
        if texts and not st.session_state.get("sum_input"):
            st.session_state["sum_input"] = ("  ".join(texts))[:200000]

    text = st.text_area("Enter article text", height=280, placeholder="Paste a long article here…", key="sum_input")
    if st.button("Summarize", type="primary", key="sum_btn"):
        if not text or not text.strip():
            st.warning("Please paste some text or upload a file.")
        else:
            with st.spinner("Loading model and generating…"):
                tok_mdl = load_model(model_name)
                if not tok_mdl or tok_mdl[0] is None:
                    st.error("Transformers not installed or failed to load model.")
                    return
                tok, mdl, device = tok_mdl
                t0 = time.time()
                out = summarize(text.strip(), tok, mdl, device, max_input_tokens=max_input_tokens, summary_max_tokens=summary_max_tokens, use_chunking=use_chunking, num_beams=num_beams)
                dt = time.time() - t0
            st.success(f"Done in {dt:.2f}s")
            st.subheader("Summary")
            st.write(out)


# ===============================
# Question Answering Page
# ===============================
def qa_page():
    st.header("❓ Question Answering")
    st.write("Ask questions given a context using Hugging Face Transformers pipelines.")

    models = [
        "distilbert-base-uncased-distilled-squad",
        "ktrapeznikov/albert-xlarge-v2-squad-v2",
    ]

    @apply_cache
    def load_qa(model_name: str):
        try:
            import torch  # type: ignore
            from transformers import pipeline  # type: ignore
        except Exception as e:
            return None, f"Transformers not available: {e}"
        device_index = 0 if hasattr(torch, "cuda") and torch.cuda.is_available() else -1
        try:
            pipe = pipeline('question-answering', model=model_name, tokenizer=model_name, device=device_index)
            return pipe, None
        except Exception as e:
            return None, str(e)

    model_name = st.sidebar.selectbox('QA model', models, index=0)
    context = st.text_area('Context', height=280)
    question = st.text_input('Question')
    if st.button('Answer', type="primary"):
        pipe, err = load_qa(model_name)
        if err or pipe is None:
            st.error(f"Failed to load QA model: {err}")
            return
        try:
            out = pipe(question=question, context=context)
            st.write('Answer:', out.get('answer'))  
        except Exception as e:
            st.error(f"Inference failed: {e}")


# ===============================
# Named Entity Recognition Page
# ===============================
def ner_page():
    st.header("🏷️ Named Entity Recognition")
    st.write("Compare spaCy NER with a simple rule-based EntityRuler.")

    CANDIDATES = ["en_core_web_md", "en_core_web_sm"]

    def build_entity_ruler(nlp, overwrite_ents=False):
        if "entity_ruler_custom" in nlp.pipe_names:
            nlp.remove_pipe("entity_ruler_custom")
        ruler = nlp.add_pipe("entity_ruler", name="entity_ruler_custom", last=True, config={"overwrite_ents": overwrite_ents})
        patterns = []
        org_suffixes = ["Inc","Inc.","Corp","Corp.","Corporation","LLC","Ltd","Ltd.","University","Institute","Group","Co.","Company"]
        for suf in org_suffixes:
            patterns.append({"label": "ORG", "pattern": [{"IS_TITLE": True, "OP": "+"}, {"TEXT": suf}]})
        titles = ["Mr.","Mrs.","Ms.","Dr.","Prof.","President","Prime Minister","Sir"]
        for t in titles:
            patterns.append({"label": "PERSON", "pattern": [{"TEXT": t}, {"IS_TITLE": True, "OP": "+"}]})
        gpes = ["United States","U.S.","USA","United Kingdom","UK","New York","London","Paris","Berlin","Tokyo"]
        for g in gpes:
            patterns.append({"label": "GPE", "pattern": g})
        orgs = ["United Nations","European Union","Apple Inc","Google","Microsoft","Facebook","Twitter"]
        for o in orgs:
            patterns.append({"label": "ORG", "pattern": o})
        ruler.add_patterns(patterns)
        return ruler

    @apply_cache
    def load_models():
        loaded = {}
        try:
            import spacy  # type: ignore
        except Exception as e:
            return {"rule_based": None}, f"spaCy not available: {e}"
        # Try to load candidates
        for name in CANDIDATES:
            try:
                loaded[name] = spacy.load(name)
            except Exception:
                pass
        base = None
        try:
            if loaded:
                base = next(iter(loaded.values()))
            else:
                base = spacy.blank("en")
        except Exception:
            base = spacy.blank("en")
        # Build rule-based
        try:
            ruler_model = base.copy()
        except Exception:
            try:
                ruler_model = spacy.blank("en")
            except Exception:
                ruler_model = None
        if ruler_model is not None:
            build_entity_ruler(ruler_model, overwrite_ents=False)
            loaded = {"rule_based": ruler_model, **loaded}
        else:
            loaded = {"rule_based": None, **loaded}
        return loaded, None

    models, err = load_models()
    if err:
        st.error(err)
        return
    names = list(models.keys())
    col1, col2 = st.columns([3, 1])
    with col2:
        model_choice = st.selectbox("Model", names, index=0)
        max_chars = st.slider("Max text length", 100, 5000, 1000, 100)
    with col1:
        default_text = "Apple Inc. hired Dr. John Smith in New York to lead AI research for the European Union."
        text = st.text_area("Input text", default_text, height=200, max_chars=max_chars, key="ner_text")
    if st.button("Analyze", key="ner_btn"):
        nlp = models.get(model_choice)
        if nlp is None:
            st.error("Selected model not available.")
            return
        doc = nlp(text)
        ents = [
            {"text": ent.text, "label": ent.label_, "start": ent.start_char, "end": ent.end_char}
            for ent in getattr(doc, "ents", [])
        ]
        df = pd.DataFrame(ents)
        if df.empty:
            st.info("No entities found.")
        else:
            labels = sorted(df["label"].unique().tolist())
            selected = st.multiselect("Filter labels", labels, default=labels)
            st.subheader("Entities")
            st.dataframe(df[df["label"].isin(selected)], use_container_width=True)
        try:
            from spacy import displacy  # type: ignore
            html = displacy.render(doc, style="ent", options={"compact": True})
            st.subheader("Highlighted text")
            st.components.v1.html(html, height=300, scrolling=True)
        except Exception:
            pass


# ===============================
# Resume Screening Page
# ===============================
def resume_screening_page():
    st.header("📄 Resume Screening")
    st.write("Rank uploaded resumes against a selected job description with semantic similarity.")

    @apply_cache
    def load_embedder():
        try:
            from sentence_transformers import SentenceTransformer  # type: ignore
        except Exception as e:
            return None, f"sentence-transformers not available: {e}"
        try:
            model = SentenceTransformer("all-MiniLM-L6-v2")
            return model, None
        except Exception as e:
            return None, str(e)

    model, err = load_embedder()
    if err:
        st.error(err)
        return

    # Prefer datasets/resumejobs/job_descriptions.csv, else fallback to jobs.csv
    def load_jobs_df():
        ds_path = os.path.join('datasets','resumejobs','job_descriptions.csv')
        if os.path.exists(ds_path):
            df = pd.read_csv(ds_path, nrows=5000)  # cap for snappy UI
        elif os.path.exists('jobs.csv'):
            df = pd.read_csv('jobs.csv')
        else:
            return None
        lower_map = {c.lower(): c for c in df.columns}
        def pick(keys):
            for k in keys:
                if k in lower_map:
                    return lower_map[k]
            return None
        desc_col = pick(['job description','description','job_description','jd','desc','text','summary','requirements','qualifications']) or df.columns[0]
        title_col = pick(['job title','title','job_title','position','role'])
        id_col = pick(['job id','job_id','jobid','id'])
        rename_map = {desc_col: 'description'}
        if title_col: rename_map[title_col] = 'title'
        if id_col: rename_map[id_col] = 'job_id'
        df = df.rename(columns=rename_map)
        if 'title' not in df.columns:
            df['title'] = df.index.astype(str)
        if 'job_id' not in df.columns:
            df['job_id'] = df.index.astype(str)
        df['description_clean'] = df['description'].astype(str).apply(normalize_text)
        return df

    jobs = load_jobs_df()
    if jobs is not None and len(jobs) > 0:
        job_idx = st.sidebar.selectbox(
            'Choose job',
            jobs.index.tolist(),
            format_func=lambda i: f"{jobs.loc[i,'title']} (id={jobs.loc[i,'job_id']})" if 'job_id' in jobs.columns else f"{jobs.loc[i,'title']} (row {i})"
        )
        job_text = jobs.loc[job_idx, 'description_clean']
    else:
        st.sidebar.info('Place datasets/resumejobs/job_descriptions.csv or jobs.csv in the app folder.')
        job_text = st.text_area('Or paste a job description here:')

    # Sidebar options
    st.sidebar.header('Options')
    top_k = st.sidebar.slider('Top-K resumes to display', 1, 50, 10)
    show_just = st.sidebar.checkbox('Show justifications for top matches', value=False)
    just_n = st.sidebar.slider('Sentences per resume (when showing justification)', 1, 5, 3)

    # Readers for uploaded files (txt, pdf, docx)
    def read_txt_file(uploaded):
        data = uploaded.read()
        try:
            return data.decode('utf-8', errors='ignore') if isinstance(data, (bytes, bytearray)) else str(data)
        finally:
            try:
                uploaded.seek(0)
            except Exception:
                pass

    def read_pdf_file(uploaded):
        text = ''
        try:
            import pdfplumber  # type: ignore
            with pdfplumber.open(io.BytesIO(uploaded.read())) as pdf:
                for p in pdf.pages:
                    text += p.extract_text() or ''
        except Exception:
            try:
                uploaded.seek(0)
            except Exception:
                pass
        finally:
            try:
                uploaded.seek(0)
            except Exception:
                pass
        return text

    def read_docx_file(uploaded):
        text = ''
        try:
            import docx  # type: ignore
            doc = docx.Document(io.BytesIO(uploaded.read()))
            text = '\n'.join([p.text for p in doc.paragraphs])
        except Exception:
            pass
        finally:
            try:
                uploaded.seek(0)
            except Exception:
                pass
        return text

    uploaded_files = st.file_uploader('Upload resumes (.pdf, .docx, .txt) — multiple allowed', type=['pdf','docx','txt'], accept_multiple_files=True, key="rs_files")

    if not job_text:
        st.warning('Provide or select a job description first.')
        return

    if not uploaded_files:
        st.info('Upload one or more resumes to score against the selected job.')
        return

    # Inference helpers
    def get_score(job_text: str, resume_text: str) -> float:
        if not job_text or not resume_text:
            return 0.0
        job_emb = model.encode([job_text], convert_to_numpy=True)
        resume_emb = model.encode([resume_text], convert_to_numpy=True)
        from sklearn.metrics.pairwise import cosine_similarity  # type: ignore
        return float(cosine_similarity(job_emb, resume_emb).squeeze())

    def justify(job_text: str, resume_text: str, k: int = 3):
        job_sents = [s.strip() for s in re.split(r'[\n\.!?]+', job_text) if s.strip()][:200]
        res_sents = [s.strip() for s in re.split(r'[\n\.!?]+', resume_text) if s.strip()][:200]
        if not job_sents or not res_sents:
            return []
        ej = model.encode(job_sents, convert_to_numpy=True)
        er = model.encode(res_sents, convert_to_numpy=True)
        from sklearn.metrics.pairwise import cosine_similarity  # type: ignore
        sim = cosine_similarity(er, ej).max(axis=1)
        idx = np.argsort(sim)[-k:][::-1]
        return [res_sents[i] for i in idx]

    rows = []
    for uf in uploaded_files:
        name = uf.name
        ext = name.split('.')[-1].lower()
        if ext == 'txt':
            resume_text = read_txt_file(uf)
        elif ext == 'pdf':
            resume_text = read_pdf_file(uf)
        elif ext == 'docx':
            resume_text = read_docx_file(uf)
        else:
            resume_text = ''
        text_clean = normalize_text(resume_text)
        score = get_score(job_text, text_clean)
        rows.append({
            'file': name,
            'chars': len(text_clean),
            'matching score': round(score*100, 2),
            'text': text_clean,
        })

    if not rows:
        st.warning('No text extracted from the uploaded files.')
        return

    res = pd.DataFrame(rows).sort_values('matching score', ascending=False).reset_index(drop=True)
    st.subheader('Top matches')
    st.dataframe(res[['file','chars','matching score']].head(top_k), use_container_width=True)
    st.download_button('Download results CSV', res.to_csv(index=False).encode('utf-8'), file_name='resume_screening_results.csv', mime='text/csv')

    if show_just:
        st.markdown('---')
        st.subheader('Justifications')
        top_subset = res.head(top_k)
        for i, row in top_subset.iterrows():
            with st.expander(f"{row['file']} — Matching Score: {row['matching score']}%"):
                sents = justify(job_text, row['text'], k=just_n)
                if sents:
                    for s in sents:
                        st.write('- ', s)
                else:
                    st.write('No sentences found.')


# ===============================
# Router
# ===============================
if APP == "Sentiment Analysis":
    sentiment_page()
elif APP == "Summarizer":
    summarizer_page()
elif APP == "Question Answering":
    qa_page()
elif APP == "Named Entity Recognition":
    ner_page()
elif APP == "Resume Screening":
    resume_screening_page()
else:
    st.info("Select a tool from the sidebar.")
